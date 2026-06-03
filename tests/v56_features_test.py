"""v5.6 tests: bulk Word -> Markdown converter (headings, bold/italic, lists,
tables) + end-to-end through a real ToolPage.

Run:  python tests/v56_features_test.py
"""
from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

failures: list[str] = []
rep = lambda m: None  # noqa: E731


def check(name, ok, detail=""):
    print(f"[{'PASS' if ok else 'FAIL'}] {name}" + (f": {detail}" if detail else ""))
    if not ok:
        failures.append(name)


def _make_docx(path: Path) -> Path:
    import docx
    d = docx.Document()
    d.add_heading("Project Report", level=1)
    d.add_heading("Overview", level=2)
    p = d.add_paragraph("This document has ")
    p.add_run("bold").bold = True
    p.add_run(" and ")
    p.add_run("italic").italic = True
    p.add_run(" text.")
    d.add_paragraph("First bullet", style="List Bullet")
    d.add_paragraph("Second bullet", style="List Bullet")
    d.add_paragraph("Step one", style="List Number")
    t = d.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Name"; t.rows[0].cells[1].text = "Value"
    t.rows[1].cells[0].text = "Alpha"; t.rows[1].cells[1].text = "42"
    d.save(str(path))
    return path


def main() -> int:
    from PySide6.QtWidgets import QApplication
    if QApplication.instance() is None:
        QApplication([])
    from mico360.core import processors as P
    from mico360.core.tools import TOOLS_BY_ID

    check("word_to_md tool registered", "word_to_md" in TOOLS_BY_ID)
    check("word_to_md accepts .docx", ".docx" in TOOLS_BY_ID["word_to_md"].accept)

    with tempfile.TemporaryDirectory() as td:
        tmp = Path(td); out = tmp / "out"; out.mkdir()
        src = _make_docx(tmp / "report.docx")
        r = P.word_to_md(src, out, {"overwrite": True}, rep)
        check("produced a .md file", r and r[0].suffix == ".md" and r[0].exists())
        md = r[0].read_text(encoding="utf-8")
        print("---- markdown ----\n" + md + "\n------------------")
        check("H1 heading", "# Project Report" in md)
        check("H2 heading", "## Overview" in md)
        check("bold rendered", "**bold**" in md)
        check("italic rendered", "*italic*" in md)
        check("bulleted list", "- First bullet" in md and "- Second bullet" in md)
        check("numbered list", "1. Step one" in md)
        check("table header + separator",
              "| Name | Value |" in md and "| --- | --- |" in md)
        check("table data row", "| Alpha | 42 |" in md)
        check("output is valid utf-8 text", isinstance(md, str) and len(md) > 0)

        # non-.docx without LibreOffice gives a clear message (not a crash)
        from mico360.core.deps import find_libreoffice
        if not find_libreoffice():
            fake = tmp / "old.doc"; fake.write_bytes(b"\xd0\xcf not really")
            try:
                P.word_to_md(fake, out, {"overwrite": True}, rep)
                check(".doc w/o LibreOffice -> clear error", False, "no error raised")
            except P.ProcessError as exc:
                check(".doc w/o LibreOffice -> clear error",
                      "LibreOffice" in str(exc) or ".docx" in str(exc))

        # --- end-to-end through a real ToolPage -------------------------
        from PySide6.QtCore import QEventLoop, QTimer
        from mico360.config import settings
        from mico360.ui.tool_page import ToolPage
        saved = settings.output_dir
        settings.output_dir = str(out)
        page = ToolPage(TOOLS_BY_ID["word_to_md"])
        page.chk_same.setChecked(False); page.chk_overwrite.setChecked(True)
        page.out_edit.setText(str(out))
        page.add_paths([str(_make_docx(tmp / "doc2.docx"))])
        res = {}; loop = QEventLoop()
        page.start()
        page.controller.finished.connect(lambda s: (res.update(s), loop.quit()))
        QTimer.singleShot(20000, loop.quit); loop.exec()
        check("UI run: 1 ok, produced .md",
              res.get("ok") == 1 and (res.get("outputs") or [None])[0].suffix == ".md",
              f"ok={res.get('ok')}")
        settings.output_dir = saved

    print()
    if failures:
        print(f"{len(failures)} check(s) failed: {', '.join(failures)}")
        return 1
    print("All v5.6 feature checks passed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
